import tkinter as tk
from tkinter import filedialog, PhotoImage, messagebox
import openpyxl
import os
import sys
from docx import Document


class ProtocolGeneratorApp:
    # Constants for file and folder paths
    CONFIG_FILE_PATH = 'Szablony Protokołów pomiarowych/Config_File.xlsx'
    CONFIG_FILE_PATH_PERSONS = 'Szablony Protokołów pomiarowych/lista_osob'
    TEMPLATES_FOLDER = 'Szablony Protokołów pomiarowych'
    LOGO_PATH = 'Szablony Protokołów pomiarowych/Logo.png'

    def __init__(self, root):
        # Initialize main application window and set up UI
        self.root = root
        self.root.title("Measurement Protocol Generator")
        self.set_icon()
        self.data = self.load_excel_data()
        if self.data:
            self.checkbox_vars = []
            self.entry_fields = []
            self.create_ui()
            self.center_window()
        else:
            # Show error if configuration data can't be loaded and close app
            messagebox.showerror("Error", "Could not load data from the configuration file.")
            self.root.destroy()

    def resource_path(self, relative_path):
        # Get absolute path to resource, works for PyInstaller and normal script execution
        try:
            base_path = sys._MEIPASS  # PyInstaller temporary folder
        except Exception:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)

    def set_icon(self):
        # Set window icon from logo file
        logo_path = self.resource_path(self.LOGO_PATH)
        try:
            logo = PhotoImage(file=logo_path)
            self.root.iconphoto(False, logo)
        except tk.TclError:
            print(f"Could not load icon from: {logo_path}")

    def load_excel_data(self):
        # Load Excel configuration file and return list of dicts representing each row
        excel_path = self.resource_path(self.CONFIG_FILE_PATH)
        try:
            workbook = openpyxl.load_workbook(excel_path)
            sheet = workbook.active
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                return []
            headers = rows[0]
            data = [dict(zip(headers, row)) for row in rows[1:]]
            return data
        except FileNotFoundError:
            messagebox.showerror("Error", f"Configuration file not found: {excel_path}")
            return None
        except openpyxl.utils.exceptions.InvalidFileException:
            messagebox.showerror("Error", f"Invalid Excel file format: {excel_path}")
            return None

    def edit_docx(self, template_path, textbox_values, headers):
        # Open a Word document template and replace placeholders with textbox values
        try:
            doc = Document(template_path)
            # Generate placeholder names based on headers, skipping the first column
            placeholders = [f"{element}_1" for element in headers[1:]]
            for placeholder, value in zip(placeholders, textbox_values):
                # Replace placeholders in paragraphs
                for p in doc.paragraphs:
                    if placeholder in p.text:
                        for run in p.runs:
                            if placeholder in run.text:
                                # Append " °C" if placeholder relates to temperature
                                run.text = run.text.replace(placeholder, f"{value} °C" if "Temperatura" in placeholder else value)
                # Replace placeholders in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if placeholder in cell.text:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(placeholder, f"{value} °C" if "Temperatura" in placeholder else value)
            # Create default filename based on template name and some textbox values (sanitize slashes)
            original_filename = os.path.basename(template_path).replace(".docx", "") + f"_{textbox_values[1].replace('/','-').replace('\\','-')}"+ f"_{textbox_values[4].replace('/','-').replace('\\','-')}"
            # Ask user where to save the modified document
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], initialfile=original_filename)
            if save_path:
                doc.save(save_path)
            else:
                messagebox.showinfo("Cancelled", "Save path not selected.")
        except FileNotFoundError:
            messagebox.showerror("Error", f"Template file not found: {template_path}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while editing the Word file: {e}")

    def create_ui(self):
        # Create the GUI interface with labels, checkboxes, entries, and buttons
        headers = list(self.data[0].keys())
        num_rows = len(self.data)
        num_cols = 4
        # Configure grid columns to expand properly
        for i in range(num_cols + 1):
            self.root.grid_columnconfigure(i, weight=1)
        # Add header labels
        tk.Label(self.root, text=headers[0], font=("Arial", 10, "bold")).grid(row=0, column=0, padx=5, pady=5, sticky="w")
        tk.Label(self.root, text="Wybierz", font=("Arial", 10, "bold")).grid(row=0, column=1, padx=5, pady=5, sticky="w")
        # Add a label and checkbox for each item in data
        for index, item in enumerate(self.data):
            var = tk.BooleanVar(value=False)
            self.checkbox_vars.append(var)
            tk.Label(self.root, text=item[headers[0]]).grid(row=index + 1, column=0, padx=5, pady=5, sticky="w")
            tk.Checkbutton(self.root, variable=var).grid(row=index + 1, column=1, padx=10, pady=5, sticky="w")
        # Add entry fields for each header except the first one
        start_row_entries = 0
        for col_index, header in enumerate(headers[1:]):
            tk.Label(self.root, text=header, font=("Arial", 10, "bold")).grid(row=start_row_entries + 1, column=2, padx=5, pady=5, sticky="w")
            entry = tk.Entry(self.root, width=50)
            self.entry_fields.append(entry)
            entry.grid(row=start_row_entries + 1, column=3, columnspan=4, padx=5, pady=5, sticky="w")
            start_row_entries += 1
        # Add buttons for creating templates, filling empty fields, clearing fields, and loading persons data
        button_row = max(num_rows + 1, start_row_entries + 1)
        create_button_make_files = tk.Button(self.root, text="Stwórz szablony", command=self.on_create_button_click)
        create_button_make_files.grid(row=button_row, column=0, columnspan=2, padx=20, pady=5, sticky="w")
        create_button_write_to_empty_cells = tk.Button(self.root, text="Uzupełnij Puste", command=self.write_to_empty_cells)
        create_button_write_to_empty_cells.grid(row=button_row, column=3, columnspan=1, padx=20, pady=5, sticky="w")
        create_button_write_to_empty_cells = tk.Button(self.root, text="Wyczyść pola", command=self.make_cells_empty)
        create_button_write_to_empty_cells.grid(row=button_row, column=4, columnspan=1, padx=20, pady=5, sticky="w")
        create_button_write_to_empty_cells = tk.Button(self.root, text="Osoba mierząca", command=self.load_measurement_check_persons, width=20, height=1)
        create_button_write_to_empty_cells.grid(row=7, column=7, columnspan=1, padx=20, pady=5, sticky="w")
        create_button_write_to_empty_cells = tk.Button(self.root, text="Osoba sprawdzająca", command=self.load_check_mesurement_person, width=20, height=1)
        create_button_write_to_empty_cells.grid(row=9, column=7, columnspan=1, padx=20, pady=5, sticky="e")
        # Load and display logo image
        logo_path = self.resource_path(self.LOGO_PATH)
        try:
            logo = PhotoImage(file=logo_path)
            logo_label = tk.Label(self.root, image=logo)
            logo_label.image = logo  # Keep a reference to avoid garbage collection
            logo_label.grid(row=button_row, column=5, columnspan=4, padx=10, pady=10, sticky="e")
        except tk.TclError:
            print(f"Could not load logo from: {logo_path}")

    def write_to_empty_cells(self):
        # Fill any empty entry fields with the placeholder "BRAK_INFORMACJI"
        for entry in self.entry_fields:
            if not entry.get().strip():
                entry.delete(0, tk.END)
                entry.insert(0, "BRAK_INFORMACJI")

    def make_cells_empty(self):
        # Clear all entry fields after user confirmation
        if messagebox.askyesno("Potwierdzenie", "Czy na pewno chcesz wyczyścić wszystkie pola?"):
            for entry in self.entry_fields:
                if entry.get().strip():
                    entry.delete(0, tk.END)
                    entry.insert(0, "")
        else:
            messagebox.showinfo("Informacja", "Zrezygnowano z czyszczenia pól.")

    def load_file_txt(self, filepath):
        # Load a .txt file from a specified directory with a file dialog
        try:
            if not os.path.isdir(filepath):
                messagebox.showerror("Błąd", f"Podana ścieżka '{filepath}' nie jest folderem.")
                return None
            txt_files = [f for f in os.listdir(filepath) if f.endswith(".txt")]
            if not txt_files:
                messagebox.showinfo("Informacja", f"W folderze '{filepath}' nie znaleziono żadnych plików .txt.")
                return None
            root = tk.Tk()
            root.withdraw()
            file_path = filedialog.askopenfilename(initialdir=filepath, title="Wybierz plik .txt", filetypes=(("Pliki tekstowe", "*.txt"), ("Wszystkie pliki", "*.*")))
            if file_path:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read().strip()
            else:
                return None
        except Exception as e:
            messagebox.showerror("Błąd", f"Wystąpił błąd podczas otwierania pliku: {e}")
            return None

    def load_check_mesurement_person(self):
        # Load and parse data for the "checking person" from a selected text file
        filepath = self.resource_path(os.path.join(self.TEMPLATES_FOLDER, 'lista_osob'))
        selected_file = self.load_file_from_directory(filepath)
        if selected_file:
            try:
                with open(selected_file, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    lines = content.strip().split('\n')
                    content_dict = {}
                    # Parse lines into a dictionary key-value pairs separated by ';'
                    for line in lines:
                        parts = line.split(';', 1)
                        if len(parts) == 2:
                            key = parts[0].strip()
                            value = parts[1].strip()
                            content_dict[key] = value
                    messagebox.showinfo("Dane z pliku", f"Znaleziono dane osoby sprawdzającej:\n{content_dict}")
                    # Insert values into specific entry fields if keys exist
                    if "osoba" in content_dict:
                        self.entry_fields[8].delete(0, tk.END)
                        self.entry_fields[8].insert(0, content_dict["osoba"])
                    if "uprawnienia" in content_dict:
                        self.entry_fields[9].delete(0, tk.END)
                        self.entry_fields[9].insert(0, content_dict["uprawnienia"])
            except FileNotFoundError:
                messagebox.showerror("Błąd", f"Nie znaleziono pliku: {selected_file}")
            except Exception as e:
                messagebox.showerror("Błąd", f"Wystąpił błąd: {e}")

    def load_measurement_check_persons(self):
        # Load and parse data for the "measuring person" from a selected text file
        filepath = self.resource_path(os.path.join(self.TEMPLATES_FOLDER, 'lista_osob'))
        selected_file = self.load_file_from_directory(filepath)
        if selected_file:
            try:
                with open(selected_file, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    lines = content.strip().split('\n')
                    content_dict = {}
                    # Parse lines into dictionary key-value pairs separated by ';'
                    for line in lines:
                        parts = line.split(';', 1)
                        if len(parts) == 2:
                            key = parts[0].strip()
                            value = parts[1].strip()
                            content_dict[key] = value

                    messagebox.showinfo("Dane z pliku", f"Znaleziono dane osoby mierzącej:\n{content_dict}")

                    # Insert values into specific entry fields if keys exist
                    if "osoba" in content_dict:
                        self.entry_fields[6].delete(0, tk.END)
                        self.entry_fields[6].insert(0, content_dict["osoba"])
                    if "uprawnienia" in content_dict:
                        self.entry_fields[7].delete(0, tk.END)
                        self.entry_fields[7].insert(0, content_dict["uprawnienia"])
            except FileNotFoundError:
                messagebox.showerror("Błąd", f"Nie znaleziono pliku: {selected_file}")
            except Exception as e:
                messagebox.showerror("Błąd", f"Wystąpił błąd podczas odczytu pliku: {e}")
        else:
            messagebox.showinfo("Informacja", "Nie wybrano pliku z osobą mierzącą.")

    def load_file_from_directory(self, filepath):
        # Open a file dialog for the user to select a .txt file from a given directory
        try:
            if not os.path.isdir(filepath):
                messagebox.showerror("Błąd", f"Podana ścieżka '{filepath}' nie jest folderem.")
                return None

            root = tk.Tk()
            root.withdraw()  # Hide main tkinter window

            file_path = filedialog.askopenfilename(
                initialdir=filepath,
                title="Wybierz plik .txt z osobą",
                filetypes=(("Pliki tekstowe", "*.txt"), ("Wszystkie pliki", "*.*"))
            )
            root.destroy()  # Close the hidden tkinter window after selection

            return file_path
        except Exception as e:
            messagebox.showerror("Błąd", f"Wystąpił błąd podczas otwierania okna dialogowego: {e}")
            return None

    def on_create_button_click(self):
        # Handle the creation of protocol documents based on selected templates and entered data

        headers = list(self.data[0].keys())
        current_directory = os.path.dirname(os.path.abspath(__file__))
        templates_folder = os.path.join(current_directory, self.TEMPLATES_FOLDER)
        textbox_values = [entry.get() for entry in self.entry_fields]

        # For each checked template checkbox, open the corresponding template and fill data
        for index, var in enumerate(self.checkbox_vars):
            if var.get():
                template_name = self.data[index].get('Rodzaj szablonu prokotołu pomiarowego')
                if template_name:
                    template_path = os.path.join(templates_folder, f"{template_name}.docx")
                    if os.path.exists(template_path):
                        self.edit_docx(template_path, textbox_values, headers)
                    else:
                        messagebox.showerror("Error", f"Template not found: {template_path}")
                else:
                    messagebox.showerror("Error", f"No template name for position {index + 1}.")

    def center_window(self):
        # Center the main application window on the user's screen
        self.root.update_idletasks()
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = self.root.winfo_width()
        window_height = self.root.winfo_height()
        position_top = int((screen_height / 2) - (window_height / 2))
        position_left = int((screen_width / 2) - (window_width / 2))
        self.root.geometry(f'{window_width}x{window_height}+{position_left}+{position_top}')


def main():
    # Entry point of the program: create the app window and run the main loop
    root = tk.Tk()
    app = ProtocolGeneratorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
